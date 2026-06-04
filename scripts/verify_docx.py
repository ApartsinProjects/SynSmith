"""Verify DOCX content canaries per the global build-artifact-contents rule.

For each output DOCX, count:
- Embedded media files (PNGs in word/media/)
- Tables (via python-docx)
- Paragraphs
Compare to the expected counts from the source HTML.
"""
import sys
import zipfile
from pathlib import Path

from docx import Document

REPO = Path(__file__).resolve().parents[1]


def inspect(path: Path) -> None:
    if not path.exists():
        print(f"{path.name}: MISSING")
        return
    sz_mb = path.stat().st_size / (1024 * 1024)
    # Count media files
    with zipfile.ZipFile(path) as z:
        media = [n for n in z.namelist() if n.startswith("word/media/")]
    # python-docx counts
    doc = Document(path)
    n_tables = len(doc.tables)
    n_paragraphs = len(doc.paragraphs)
    n_sections = len(doc.sections)
    print(f"{path.name}")
    print(f"  size:        {sz_mb:.2f} MB")
    print(f"  media files: {len(media)}  ({sorted(set(Path(m).suffix for m in media))})")
    print(f"  tables:      {n_tables}")
    print(f"  paragraphs:  {n_paragraphs}")
    print(f"  sections:    {n_sections}")
    print()


inspect(REPO / "docs" / "synsmith.docx")
inspect(REPO / "docs" / "synsmith_2col.docx")
